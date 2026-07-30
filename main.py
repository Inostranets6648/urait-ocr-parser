import asyncio
import os
import random
import re
from pathlib import Path
from docx import Document
import fitz
from PIL import Image
from playwright.async_api import async_playwright
import pytesseract

BASE_DIR = Path(__file__).resolve().parent

def setup_tesseract():
    """Автоматически ищет Tesseract OCR в системе или просит путь у пользователя."""
    default_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.expanduser(r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe")
    ]
    
    for path in default_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            return True
            
    user_path = input("Tesseract OCR не найден автоматически. Укажите путь к tesseract.exe: ").strip('"')
    if os.path.exists(user_path):
        pytesseract.pytesseract.tesseract_cmd = user_path
        return True
    else:
        print("[ОШИБКА] Указанный путь к Tesseract не существует!")
        return False

START_PARSING = False

def convert_svg_to_png(svg_data: bytes, png_path: Path, scale=3.0) -> bool:
    try:
        doc = fitz.open(stream=svg_data, filetype="svg")
        page = doc[0]
        mat = fitz.Matrix(scale, scale)
        pix = page.get_pixmap(matrix=mat)
        pix.save(str(png_path))
        return True
    except Exception as e:
        print(f"Ошибка конвертации SVG: {e}")
        return False

def ocr_and_clean(png_path: Path) -> str:
    try:
        raw_text = pytesseract.image_to_string(Image.open(png_path), lang="rus+eng")
        text = re.sub(r"(\w+)\s*[-—:\u00AD]\s*\n\s*(\w+)", r"\1\2", raw_text)
        text = re.sub(r"\n\s*\n", "[[PARAGRAPH]]", text)
        text = re.sub(r"\n", " ", text)
        text = text.replace("[[PARAGRAPH]]", "\n\n")
        text = re.sub(r" +", " ", text)
        return text.strip()
    except Exception as e:
        print(f"Ошибка OCR: {e}")
        return ""

def append_page_to_files(page_num: int, text: str, txt_path: Path, docx_path: Path):
    """Постранично дописывает результат в файлы на диске."""
    with open(txt_path, "a", encoding="utf-8") as f:
        f.write(f"\n\n=== СТРАНИЦА {page_num} ===\n\n")
        f.write(text)

    if not docx_path.exists():
        doc = Document()
    else:
        doc = Document(str(docx_path))

    doc.add_heading(f"Страница {page_num}", level=2)
    for paragraph in text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    doc.add_page_break()
    doc.save(str(docx_path))

def convert_docx_to_pdf(docx_path: Path, pdf_path: Path):
    """Генерирует финальный PDF с поддержкой кириллицы из готового DOCX."""
    if not docx_path.exists():
        return

    print("\n[СБОРКА] Формирование итогового PDF файла...")
    doc = Document(str(docx_path))
    pdf_doc = fitz.open()

    current_page_text = []
    page_counter = 1

    for p in doc.paragraphs:
        if p.text.startswith("Страница "):
            if current_page_text:
                pdf_page = pdf_doc.new_page(width=595, height=842)
                html_content = f"<h3>Страница {page_counter}</h3>" + "".join(
                    f"<p>{txt}</p>" for txt in current_page_text
                )
                pdf_page.insert_htmlbox(fitz.Rect(40, 40, 555, 802), html_content)
                current_page_text = []
                page_counter += 1
        else:
            if p.text.strip():
                current_page_text.append(p.text.strip())

    if current_page_text:
        pdf_page = pdf_doc.new_page(width=595, height=842)
        html_content = f"<h3>Страница {page_counter}</h3>" + "".join(
            f"<p>{txt}</p>" for txt in current_page_text
        )
        pdf_page.insert_htmlbox(fitz.Rect(40, 40, 555, 802), html_content)

    pdf_doc.save(str(pdf_path))
    print(f"[УСПЕХ] Готовый PDF сохранен в: {pdf_path}")

async def main():
    global START_PARSING

    if not setup_tesseract():
        return

    print("=== Web-Reader OCR Extractor ===")
    target_url = input("Вставьте URL книги/документа: ").strip()
    
    folder_input = input("Папка для сохранения (нажмите Enter для 'output'): ").strip()
    output_dir = BASE_DIR / (folder_input if folder_input else "output")
    output_dir.mkdir(parents=True, exist_ok=True)

    max_pages = int(input("Сколько страниц выгрузить? (по умолчанию 50): ") or "50")

    txt_path = output_dir / "extracted_book.txt"
    docx_path = output_dir / "extracted_book.docx"
    pdf_path = output_dir / "extracted_book.pdf"

    for p in [txt_path, docx_path, pdf_path]:
        if p.exists():
            os.remove(p)

    seen_svg_hashes = set()
    page_counter = 0

    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=False)
        context = await browser.new_context()
        page = await context.new_page()

        async def handle_response(response):
            nonlocal page_counter
            if not START_PARSING:
                return

            url = response.url
            if ".svg" in url or "page" in url or "viewer" in url:
                try:
                    body = await response.body()

                    if len(body) > 15000 and b"<svg" in body[:200] and b"font_" in body:
                        body_hash = hash(body)
                        if body_hash in seen_svg_hashes:
                            return
                        seen_svg_hashes.add(body_hash)

                        page_counter += 1
                        current_page = page_counter
                        print(f"\n[Страница {current_page}] Захват ({len(body)} байт)...")

                        temp_png = output_dir / f"temp_{current_page}.png"

                        if convert_svg_to_png(body, temp_png):
                            text = ocr_and_clean(temp_png)

                            if len(text) > 50:
                                append_page_to_files(current_page, text, txt_path, docx_path)
                                print(f"Страница {current_page} сохранена на диск ({len(text)} символов).")

                            if temp_png.exists():
                                os.remove(temp_png)
                except Exception:
                    pass

        page.on("response", handle_response)

        print("\nЗапуск браузера...")
        await page.goto(target_url)

        print("\n--- ИНСТРУКЦИЯ ---")
        print("1. Войдите в систему (если требуется).")
        print("2. Перейдите на страницу, с которой нужно начать выгрузку.")
        input("\nНажмите ENTER в этой консоли, чтобы начать автоматический парсинг...")

        START_PARSING = True
        print(f"\nНачато автоматическое пролистывание ({max_pages} страниц)...\n")

        for i in range(max_pages):
            print(f"Листаем дальше ({i+1}/{max_pages})...", end="\r")
            await page.keyboard.press("ArrowRight")
            await asyncio.sleep(3.5 + random.uniform(0.2, 0.6))

        await browser.close()

    print("\nВыгрузка полностью завершена!")
    convert_docx_to_pdf(docx_path, pdf_path)

if __name__ == "__main__":
    asyncio.run(main())
